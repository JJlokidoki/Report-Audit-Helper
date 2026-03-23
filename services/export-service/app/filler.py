import re
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from docxtpl import DocxTemplate
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.html_to_docx import enrich_context, enrich_context_with_subdoc

_DRAWING_NS = {
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}


def _patch_missing_ns(buf: BytesIO) -> BytesIO:
    """Добавляет недостающие xmlns-декларации для drawing-элементов из Subdoc."""
    buf.seek(0)
    with ZipFile(buf, "r") as zin:
        xml = zin.read("word/document.xml").decode("utf-8")

        patches = [
            f'xmlns:{p}="{u}"'
            for p, u in _DRAWING_NS.items()
            if f"{p}:" in xml and f"xmlns:{p}=" not in xml
        ]
        if not patches:
            buf.seek(0)
            return buf

        xml = xml.replace(
            "<w:document ", "<w:document " + " ".join(patches) + " ", 1
        )

        files = {i.filename: zin.read(i.filename) for i in zin.infolist()}

    out = BytesIO()
    with ZipFile(out, "w") as zout:
        for name, data in files.items():
            zout.writestr(name, xml.encode("utf-8") if name == "word/document.xml" else data)
    out.seek(0)
    return out


def _merge_jinja_runs(template_path: Path) -> BytesIO:
    """Склеивает разбитые Word-ом runs содержащие Jinja-теги в один run.

    Word часто разбивает текст вроде {%tr endfor %} на несколько XML-runs
    с разным форматированием, из-за чего docxtpl не может распознать тег.
    """
    buf = BytesIO()
    with ZipFile(str(template_path), "r") as zin:
        xml = zin.read("word/document.xml").decode("utf-8")
        # Склеиваем соседние <w:r>...</w:r> внутри <w:p> и <w:tc>,
        # если их объединённый текст содержит Jinja-маркеры
        def _merge_runs_in_match(m: re.Match) -> str:
            tag = m.group(0)
            # Находим все <w:r>...</w:r> блоки
            runs = list(re.finditer(r"<w:r[ >].*?</w:r>", tag, re.DOTALL))
            if len(runs) < 2:
                return tag
            # Собираем текст из всех runs
            full_text = ""
            for run in runs:
                texts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", run.group(0), re.DOTALL)
                full_text += "".join(texts)
            # Если нет Jinja-тегов, не трогаем
            if not re.search(r"\{[%\{]", full_text):
                return tag
            # Берём форматирование из первого run
            first_run = runs[0].group(0)
            rpr_match = re.search(r"<w:rPr>.*?</w:rPr>", first_run, re.DOTALL)
            rpr = rpr_match.group(0) if rpr_match else ""
            # Собираем один run со всем текстом
            merged_run = f'<w:r>{rpr}<w:t xml:space="preserve">{full_text}</w:t></w:r>'
            # Заменяем все runs на один
            start = runs[0].start()
            end = runs[-1].end()
            return tag[:start] + merged_run + tag[end:]

        # Обрабатываем ячейки таблицы и параграфы
        xml = re.sub(r"<w:tc[ >].*?</w:tc>", _merge_runs_in_match, xml, flags=re.DOTALL)
        xml = re.sub(r"<w:p[ >].*?</w:p>", _merge_runs_in_match, xml, flags=re.DOTALL)

        files = {i.filename: zin.read(i.filename) for i in zin.infolist()}

    with ZipFile(buf, "w") as zout:
        for name, data in files.items():
            zout.writestr(name, xml.encode("utf-8") if name == "word/document.xml" else data)
    buf.seek(0)
    return buf


def _find_table_by_marker(doc: Document, marker: str) -> tuple[int, int] | None:
    """Ищет таблицу и строку содержащую маркер. Возвращает (table_index, row_index)."""
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for cell in row.cells:
                if marker in cell.text:
                    return ti, ri
    return None


def _fill_table_rows(doc_buf: BytesIO, marker: str,
                     rows_data: list[list[str]]) -> BytesIO:
    """Заполняет таблицу по маркеру.

    Находит таблицу содержащую marker (например «__software__»),
    клонирует строку с маркером для каждого элемента rows_data.
    """
    from copy import deepcopy

    doc = Document(doc_buf)
    found = _find_table_by_marker(doc, marker)
    if not found or not rows_data:
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    table_index, template_row = found
    table = doc.tables[table_index]
    template_row_el = table.rows[template_row]._tr

    for i, values in enumerate(rows_data):
        if i == 0:
            row = table.rows[template_row]
        else:
            new_tr = deepcopy(template_row_el)
            table._tbl.append(new_tr)
            row = table.rows[-1]

        for ci, val in enumerate(values):
            if ci < len(row.cells):
                p = row.cells[ci].paragraphs[0]
                p.text = ""
                p.add_run(val)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _make_group_header_row(table, template_row_el, group_name: str, col_count: int):
    """Создаёт строку-заголовок группы с объединёнными ячейками и жирным текстом по центру."""
    from copy import deepcopy

    new_tr = deepcopy(template_row_el)
    table._tbl.append(new_tr)
    row = table.rows[-1]

    # Очищаем все ячейки перед merge (убираем лишние параграфы)
    for ci in range(col_count):
        cell = row.cells[ci]
        for para in cell.paragraphs:
            para.text = ""

    # Объединяем все ячейки в строке
    if col_count > 1:
        row.cells[0].merge(row.cells[col_count - 1])

    # Заполняем объединённую ячейку: жирный текст по центру
    merged_cell = row.cells[0]
    # Удаляем лишние пустые параграфы, оставшиеся после merge
    while len(merged_cell.paragraphs) > 1:
        last_p = merged_cell.paragraphs[-1]._element
        last_p.getparent().remove(last_p)
    p = merged_cell.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(group_name)
    run.bold = True


def _fill_checklist_table(doc_buf: BytesIO, checks: list[dict]) -> BytesIO:
    """Заполняет таблицу чеклиста программно, клонируя строку-шаблон.

    Поддерживает группировку: если category меняется между проверками,
    вставляется строка-заголовок группы (объединённые ячейки, жирный шрифт).
    """
    from copy import deepcopy
    from docx.enum.section import WD_ORIENT

    doc = Document(doc_buf)

    # Альбомная ориентация для чеклиста
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        if section.page_width < section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width

    if not doc.tables or not checks:
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    table = doc.tables[0]
    # Сохраняем копию шаблонной строки ДО любых модификаций
    template_row_el = deepcopy(table.rows[1]._tr)
    col_count = len(table.columns)

    # Определяем нужны ли заголовки групп (больше одной категории)
    categories = list(dict.fromkeys(c.get("category", "") for c in checks))
    use_group_headers = len(categories) > 1

    # Удаляем оригинальную шаблонную строку — будем добавлять через deepcopy
    table._tbl.remove(table.rows[1]._tr)

    current_category = None

    for check in checks:
        category = check.get("category", "")

        # Вставляем заголовок группы при смене категории
        if use_group_headers and category != current_category:
            current_category = category
            _make_group_header_row(table, template_row_el, category, col_count)

        # Вставляем строку проверки
        new_tr = deepcopy(template_row_el)
        table._tbl.append(new_tr)
        row = table.rows[-1]

        values = [
            check.get("check_id", ""),
            check.get("name", ""),
            check.get("category", ""),
            check.get("result", ""),
        ]
        for ci, val in enumerate(values):
            if ci < len(row.cells):
                cell = row.cells[ci]
                p = cell.paragraphs[0]
                p.text = ""
                run = p.add_run(val)
                if ci == len(values) - 1:
                    run.bold = True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def fill_template(template_path: Path, context: dict, use_subdoc: bool = False) -> BytesIO:
    """Заполняет docxtpl шаблон контекстом, конвертируя HTML-поля в RichText."""
    cleaned = _merge_jinja_runs(template_path)
    doc = DocxTemplate(cleaned)
    enrich = enrich_context_with_subdoc if use_subdoc else enrich_context

    # Извлекаем данные для программного заполнения таблиц — они не нужны Jinja
    checks = context.pop("checks", None)
    software = context.pop("software", None)

    doc.render(enrich(context, doc))
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    buf = _patch_missing_ns(buf)

    # Программное заполнение таблиц по маркерам
    if software:
        rows = [[s.get("name", ""), s.get("description", "") or ""] for s in software]
        buf = _fill_table_rows(buf, "__software__", rows)


    if checks:
        buf = _fill_checklist_table(buf, checks)

    return buf
