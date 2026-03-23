import logging
from io import BytesIO
from pathlib import Path

import httpx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxcompose.composer import Composer

from app.config import settings
from app.filler import fill_template

logger = logging.getLogger(__name__)

SEVERITY_ORDER = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}

STATUS_LABEL = {
    "passed": "Выполнено",
    "failed": "Сломано",
    "not_applicable": "Не применимо",
    "not_tested": "Не выполнено",
}

# Порядок шаблонов и их номера глав (None — без нумерации: титул, оглавление)
TEMPLATE_FILES: list[tuple[str, int | None]] = [
    ("01_title.docx",              None),
    ("02_toc.docx",                None),
    ("03_general_info.docx",       1),
    ("04_test_results.docx",       2),
    # 05_vulnerability.docx вставляется динамически как глава 3
    ("06_threat_classification.docx", 4),
    ("07_checklist.docx",          5),
]

VULN_TEMPLATE = "05_vulnerability.docx"
VULN_CHAPTER = 3


class _Counter:
    """Сквозной счётчик для нумерации таблиц/рисунков между шаблонами."""

    def __init__(self, start: int = 0):
        self._value = start

    def __call__(self) -> int:
        self._value += 1
        return self._value

    @property
    def current(self) -> int:
        return self._value


async def _fetch(client: httpx.AsyncClient, path: str) -> dict:
    resp = await client.get(f"{settings.report_service_url}/api{path}")
    resp.raise_for_status()
    logger.debug("Fetched %s -> %d", path, resp.status_code)
    return resp.json()


def _build_contexts(report: dict, system_info: dict, summary: dict, checklist: list, vulnerabilities: list | None = None) -> dict[str, dict]:
    """Собирает контексты для каждого шаблона."""
    executors_str = "\n".join(e["name"] for e in system_info.get("executors", []))
    software_list = system_info.get("software", [])
    counts = summary.get("counts", {})

    base = {
        "report_name": report.get("name", ""),
        "report_type": report.get("report_type", ""),
        "description": system_info.get("description") or "",
        "asName": system_info.get("asName") or "",
        "keId": system_info.get("keId") or "",
        "url": system_info.get("url") or "",
        "dateStart": system_info.get("dateStart") or "",
        "dateEnd": system_info.get("dateEnd") or "",
        "segment": system_info.get("segment") or "",
        "goal": system_info.get("goal") or "",
        "qualificationLevel": system_info.get("qualificationLevel") or "",
        "accessLevel": system_info.get("accessLevel") or "",
        "knowledgeLevel": system_info.get("knowledgeLevel") or "",
        "testConditions": system_info.get("testConditions") or "",
        "executors": executors_str,
        "executors_list": system_info.get("executors", []),
        "software": software_list,
    }

    vuln_names = [v.get("bug_name", "") for v in (vulnerabilities or [])]

    test_results = {
        **base,
        "critical_count": counts.get("critical", 0),
        "high_count": counts.get("high", 0),
        "medium_count": counts.get("medium", 0),
        "low_count": counts.get("low", 0),
        "info_count": counts.get("info", 0),
        "total_count": sum(counts.values()),
        "vulnerabilities": vulnerabilities or [],
        "vuln_names": vuln_names,
    }

    def _check_result(check: dict) -> str:
        label = STATUS_LABEL.get(check.get("status", ""), check.get("status", ""))
        notes = (check.get("notes") or "").strip()
        return f"{label}\n{notes}" if notes else label

    enriched_checks = [{**c, "result": _check_result(c)} for c in checklist]
    checklist_ctx = {**base, "checks": enriched_checks}

    return {
        "01_title.docx": base,
        "02_toc.docx": base,
        "03_general_info.docx": base,
        "04_test_results.docx": test_results,
        "06_threat_classification.docx": base,
        "07_checklist.docx": checklist_ctx,
    }


def _vuln_context(vuln: dict, is_first: bool = False, vuln_index: int = 1) -> dict:
    return {
        "is_first": is_first,
        "chapter_num": VULN_CHAPTER,
        "sub": lambda n, c=VULN_CHAPTER: f"{c}.{n}",
        "vuln_index": vuln_index,          # для подзаголовка «3.1 Название»
        "bug_name": vuln.get("bug_name", ""),
        "bug_criticality": vuln.get("bug_criticality", ""),
        "cvss_score": vuln.get("cvss_score") or "",
        "cvss_vector": vuln.get("cvss_vector") or "",
        "bug_description": vuln.get("bug_description") or "",
        "reproduction_steps": vuln.get("reproduction_steps") or "",
        "remediation": vuln.get("remediation") or "",
        "automation_level": vuln.get("automation_level", ""),
    }


def _prepend_page_break(doc: Document) -> None:
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    doc.element.body.insert(0, p)


def _is_landscape(doc: Document) -> bool:
    """Проверяет, имеет ли документ альбомную ориентацию."""
    for section in doc.sections:
        orient = section._sectPr.find(qn("w:pgSz"))
        if orient is not None and orient.get(qn("w:orient")) == "landscape":
            return True
    return False


def _prepend_section_break_landscape(doc: Document, prev_section) -> None:
    """Вставляет section break перед документом для перехода в landscape."""
    # Создаём параграф с минимальным sectPr (portrait) — это закрывает предыдущую секцию
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")

    sect_pr = OxmlElement("w:sectPr")
    # Тип: новая страница
    type_el = OxmlElement("w:type")
    type_el.set(qn("w:val"), "nextPage")
    sect_pr.append(type_el)
    # Размер страницы portrait (A4), EMU → twips (1 twip = 635 EMU)
    pg_sz = OxmlElement("w:pgSz")
    pg_sz.set(qn("w:w"), str(int(prev_section.page_width) // 635))
    pg_sz.set(qn("w:h"), str(int(prev_section.page_height) // 635))
    sect_pr.append(pg_sz)

    pPr.append(sect_pr)
    p.append(pPr)
    doc.element.body.insert(0, p)


def _merge_docs(docs: list[BytesIO]) -> BytesIO:
    """Сливает список BytesIO docx в один документ."""
    if not docs:
        raise ValueError("No documents to merge")

    master = Document(docs[0])
    composer = Composer(master)
    # TODO: какая то срань. потом разберусь
    # Обход docxcompose IndexError при doc.sections > master.sections
    def _wrap_safe(orig_fn):
        def _inner(doc):
            try:
                orig_fn(doc)
            except IndexError:
                pass
        return _inner
    composer.fix_section_types = _wrap_safe(composer.fix_section_types)
    composer.fix_header_and_footers = _wrap_safe(composer.fix_header_and_footers)

    landscape_indices = []
    for i, doc_buf in enumerate(docs[1:], start=1):
        sub = Document(doc_buf)
        if _is_landscape(sub):
            landscape_indices.append(i)
            _prepend_section_break_landscape(sub, master.sections[-1])
        else:
            _prepend_page_break(sub)
        composer.append(sub)

    # Сквозная нумерация страниц: убираем рестарт нумерации во всех секциях кроме первой
    for i, section in enumerate(master.sections):
        pg_num_type = section._sectPr.find(qn("w:pgNumType"))
        if pg_num_type is not None and i > 0:
            pg_num_type.attrib.pop(qn("w:start"), None)

    result = BytesIO()
    composer.save(result)
    result.seek(0)

    # Восстановить landscape-ориентацию после composer.save(),
    # т.к. docxcompose сбрасывает sectPr при слиянии
    if landscape_indices:
        result = _fix_landscape_sections(result)

    return result


def _fix_landscape_sections(buf: BytesIO) -> BytesIO:
    """Устанавливает landscape для последней секции документа.

    docxcompose может поместить sectPr не в body, а в pPr последнего параграфа.
    Здесь мы гарантируем что body-level sectPr имеет landscape ориентацию.
    """
    doc = Document(buf)
    body = doc.element.body

    # Ищем body-level sectPr (прямые потомки body, не вложенные в pPr)
    last_sect = body.findall(qn("w:sectPr"))
    if last_sect:
        sect_pr = last_sect[-1]
    else:
        # Нет body-level sectPr — создаём
        sect_pr = OxmlElement("w:sectPr")
        body.append(sect_pr)

    # Устанавливаем landscape
    pg_sz = sect_pr.find(qn("w:pgSz"))
    if pg_sz is None:
        pg_sz = OxmlElement("w:pgSz")
        sect_pr.insert(0, pg_sz)
    pg_sz.set(qn("w:w"), "16838")
    pg_sz.set(qn("w:h"), "11906")
    pg_sz.set(qn("w:orient"), "landscape")

    # Margins для landscape (если отсутствуют)
    pg_mar = sect_pr.find(qn("w:pgMar"))
    if pg_mar is None:
        pg_mar = OxmlElement("w:pgMar")
        pg_mar.set(qn("w:top"), "1021")
        pg_mar.set(qn("w:right"), "1021")
        pg_mar.set(qn("w:bottom"), "1021")
        pg_mar.set(qn("w:left"), "1021")
        pg_mar.set(qn("w:header"), "680")
        pg_mar.set(qn("w:footer"), "709")
        sect_pr.append(pg_mar)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


async def generate_report(report_id: int) -> BytesIO:
    logger.debug("Fetching report data: report_id=%d", report_id)
    async with httpx.AsyncClient(timeout=30, trust_env=False) as client:
        report = await _fetch(client, f"/reports/{report_id}")
        system_info = await _fetch(client, f"/reports/{report_id}/system-info")
        summary = await _fetch(client, f"/reports/{report_id}/test-summary")
        checklist = await _fetch(client, f"/reports/{report_id}/checklist")

    report_type = report.get("report_type", "web")
    template_dir = Path(settings.template_dir) / report_type

    vulnerabilities = sorted(
        summary.get("vulnerabilities", []),
        key=lambda v: SEVERITY_ORDER.get(v.get("bug_criticality", "info"), 99),
    )
    logger.info(
        "Generating report: id=%d type=%s vulns=%d",
        report_id, report_type, len(vulnerabilities),
    )

    contexts = _build_contexts(report, system_info, summary, checklist, vulnerabilities)
    filled_docs: list[BytesIO] = []

    next_table = _Counter()
    next_fig = _Counter()

    def _shared_ctx(extra: dict) -> dict:
        return {**extra, "next_table": next_table, "next_fig": next_fig}

    for filename, chapter_num in TEMPLATE_FILES:
        path = template_dir / filename
        if not path.exists():
            logger.debug("Template not found, skipping: %s", filename)
            continue
        ctx = {**contexts.get(filename, {}), "chapter_num": chapter_num}
        ctx["sub"] = lambda n, c=chapter_num: f"{c}.{n}" if c is not None else str(n)
        filled_docs.append(fill_template(path, _shared_ctx(ctx)))
        logger.debug("Filled template: %s", filename)

        if filename == "04_test_results.docx":
            vuln_path = template_dir / VULN_TEMPLATE
            if vuln_path.exists() and vulnerabilities:
                for i, vuln in enumerate(vulnerabilities):
                    vuln_ctx = _vuln_context(vuln, is_first=(i == 0), vuln_index=i + 1)
                    filled_docs.append(fill_template(vuln_path, _shared_ctx(vuln_ctx), use_subdoc=True))
                logger.debug("Filled %d vulnerability templates", len(vulnerabilities))

    if not filled_docs:
        raise FileNotFoundError(f"No templates found in {template_dir}")

    result = _merge_docs(filled_docs)
    logger.info("Report generated: id=%d docs=%d", report_id, len(filled_docs))
    return result
