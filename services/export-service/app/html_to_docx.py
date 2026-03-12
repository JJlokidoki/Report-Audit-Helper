import base64
import logging
import re
import tempfile
from io import BytesIO
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm
from docxtpl import DocxTemplate

logger = logging.getLogger(__name__)

MAX_IMAGE_WIDTH = Cm(14)

RICH_TEXT_FIELDS = frozenset(
    {"goal", "testConditions", "bug_description", "reproduction_steps", "remediation"}
)


def is_html(text: str | None) -> bool:
    return bool(text and re.search(r"<[a-zA-Z]", text))


def html_to_text(html: str | None) -> str:
    """Plain text из HTML: параграфы → \\n, li → «• item», img → [изображение]."""
    if not html:
        return ""
    if not is_html(html):
        return html

    soup = BeautifulSoup(html, "html.parser")
    parts: list[str] = []

    def process(node: NavigableString | Tag) -> None:
        if isinstance(node, NavigableString):
            parts.append(str(node))
        elif isinstance(node, Tag):
            name = node.name
            if name in ("p", "div", "h1", "h2", "h3", "h4"):
                if parts and not parts[-1].endswith("\n"):
                    parts.append("\n")
                for child in node.children:
                    process(child)
                if parts and not parts[-1].endswith("\n"):
                    parts.append("\n")
            elif name == "br":
                parts.append("\n")
            elif name in ("ul", "ol"):
                for child in node.children:
                    process(child)
            elif name == "li":
                if parts and not parts[-1].endswith("\n"):
                    parts.append("\n")
                parts.append("• ")
                for child in node.children:
                    process(child)
                if parts and not parts[-1].endswith("\n"):
                    parts.append("\n")
            elif name == "img":
                parts.append("[изображение]")
            else:
                for child in node.children:
                    process(child)

    for child in soup.children:
        process(child)
    return "".join(parts).strip()


def _clean_html(html: str) -> str:
    return re.sub(r"<p[^>]*>\s*</p>", "", html)


def _extract_base64_images(html: str) -> tuple[str, list[Path]]:
    """Заменяет data-URI в <img src> на временные файлы."""
    soup = BeautifulSoup(html, "html.parser")
    tmp_files: list[Path] = []
    for img in soup.find_all("img"):
        src = img.get("src", "")
        m = re.match(r"data:image/(\w+);base64,(.+)", src, re.DOTALL)
        if not m:
            continue
        ext = "jpg" if m.group(1) == "jpeg" else m.group(1)
        tmp = tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False)
        tmp.write(base64.b64decode(m.group(2)))
        tmp.close()
        tmp_files.append(Path(tmp.name))
        img["src"] = tmp.name
    return str(soup), tmp_files


_ALIGN_MAP = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _get_alignment(tag: Tag) -> WD_ALIGN_PARAGRAPH | None:
    style = tag.get("style", "")
    m = re.search(r"text-align:\s*(\w+)", style)
    return _ALIGN_MAP.get(m.group(1)) if m else None


def _add_inline(paragraph, node: Tag, **extra_fmt) -> None:
    """Рекурсивно добавляет inline-контент (bold/italic/text) в параграф."""
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                run = paragraph.add_run(text)
                if extra_fmt.get("bold"):
                    run.bold = True
                if extra_fmt.get("italic"):
                    run.italic = True
                if extra_fmt.get("underline"):
                    run.underline = True
        elif isinstance(child, Tag):
            if child.name in ("b", "strong"):
                _add_inline(paragraph, child, **{**extra_fmt, "bold": True})
            elif child.name in ("i", "em"):
                _add_inline(paragraph, child, **{**extra_fmt, "italic": True})
            elif child.name == "u":
                _add_inline(paragraph, child, **{**extra_fmt, "underline": True})
            elif child.name == "br":
                paragraph.add_run("\n")
            else:
                _add_inline(paragraph, child, **extra_fmt)


def _fix_subdoc_images(sub, tpl: DocxTemplate) -> None:
    """Копирует image-части из Subdoc в основной шаблон и обновляет rId."""
    doc = sub.subdocx
    rid_map: dict[str, str] = {}
    package = tpl.docx.part.package
    for rel in list(doc.part.rels.values()):
        if "image" in rel.reltype:
            image_part = package.get_or_add_image_part(
                BytesIO(rel.target_part.blob)
            )
            new_rid = tpl.docx.part.relate_to(image_part, rel.reltype)
            rid_map[rel.rId] = new_rid
    if not rid_map:
        return
    logger.debug("Migrated %d image(s) from subdoc to main template", len(rid_map))
    for blip in doc.element.body.iter(qn("a:blip")):
        old = blip.get(qn("r:embed"))
        if old in rid_map:
            blip.set(qn("r:embed"), rid_map[old])


def _add_block_image(
    doc, img_tag: Tag, parent_align: WD_ALIGN_PARAGRAPH | None = None
) -> None:
    """Вставляет <img> как отдельный параграф с выравниванием."""
    src = img_tag.get("src", "")
    if not (src and Path(src).exists()):
        logger.warning("Image not found, skipping: %s", src[:80] if src else "(empty)")
        return
    width = _get_capped_width(src)
    doc.add_picture(src, width=width)
    doc.paragraphs[-1].alignment = (
        parent_align or _get_alignment(img_tag) or WD_ALIGN_PARAGRAPH.CENTER
    )


def _get_capped_width(image_path: str) -> Cm | None:
    """Ширина для изображения, ограниченная MAX_IMAGE_WIDTH."""
    try:
        from docx.image.image import Image
        native_w = Image.from_file(image_path).native_size[0]
        return MAX_IMAGE_WIDTH if native_w > MAX_IMAGE_WIDTH else None
    except Exception:
        return MAX_IMAGE_WIDTH


def html_to_subdoc(html: str | None, tpl: DocxTemplate):
    """
    Конвертирует TipTap HTML → docxtpl Subdoc.
    Строит документ напрямую через python-docx (без htmldocx).
    Используется в шаблонах через {{p field_doc }}.
    """
    if not html:
        return ""
    if not is_html(html):
        return html

    cleaned = _clean_html(html)
    prepared, tmp_files = _extract_base64_images(cleaned)
    try:
        soup = BeautifulSoup(prepared, "html.parser")
        sub = tpl.new_subdoc()
        doc = sub.subdocx

        # Удаляем дефолтный пустой параграф
        if doc.paragraphs and not doc.paragraphs[0].text:
            p_elem = doc.paragraphs[0]._element
            p_elem.getparent().remove(p_elem)

        def process_block(elem: Tag | NavigableString) -> None:
            if isinstance(elem, NavigableString):
                text = str(elem).strip()
                if text:
                    doc.add_paragraph(text)
                return
            if not isinstance(elem, Tag):
                return

            if elem.name in ("p", "div", "h1", "h2", "h3", "h4"):
                align = _get_alignment(elem)
                inner_img = elem.find("img")
                if inner_img:
                    _add_block_image(doc, inner_img, align)
                else:
                    p = doc.add_paragraph()
                    p.alignment = align
                    _add_inline(p, elem)
            elif elem.name == "img":
                _add_block_image(doc, elem)
            elif elem.name in ("ul", "ol"):
                align = _get_alignment(elem)
                for li in elem.find_all("li", recursive=False):
                    p = doc.add_paragraph()
                    p.alignment = _get_alignment(li) or align
                    p.add_run("• ")
                    _add_inline(p, li)
            else:
                for child in elem.children:
                    process_block(child)

        for child in soup.children:
            process_block(child)

        _fix_subdoc_images(sub, tpl)
        return sub
    finally:
        for p in tmp_files:
            p.unlink(missing_ok=True)


def enrich_context(context: dict, tpl: DocxTemplate) -> dict:
    """
    Конвертирует HTML-строки в plain text для {{ field }} плейсхолдеров.
    Subdoc не создаётся — это предотвращает конфликты секций в docxcompose.
    Для шаблонов с {{p field_doc }} используй enrich_context_with_subdoc().
    """
    result: dict = {}
    for key, value in context.items():
        if key in RICH_TEXT_FIELDS and isinstance(value, str):
            result[key] = html_to_text(value)
        else:
            result[key] = value
    return result


def enrich_context_with_subdoc(context: dict, tpl: DocxTemplate) -> dict:
    """
    Расширенный вариант: дополнительно создаёт Subdoc для {{p field_doc }}.
    Используй только в шаблонах, где гарантированно нет merge через docxcompose,
    или когда шаблон явно использует {{p field_doc }} синтаксис.
    """
    result = enrich_context(context, tpl)
    for key in list(result.keys()):
        if key in RICH_TEXT_FIELDS:
            original = context.get(key)
            if isinstance(original, str):
                result[f"{key}_doc"] = html_to_subdoc(original, tpl)
    return result
