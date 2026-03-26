from io import BytesIO

from docx import Document

from app.extractors import extract_from_docx


def _make_docx(paragraphs: list[str]) -> bytes:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_from_docx_basic():
    content = _make_docx(["Hello World", "Second paragraph"])
    result = extract_from_docx(content)
    assert "Hello World" in result
    assert "Second paragraph" in result


def test_extract_from_docx_empty():
    content = _make_docx([])
    result = extract_from_docx(content)
    assert result.strip() == ""


def test_extract_from_docx_with_table():
    doc = Document()
    doc.add_paragraph("Before table")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cell 1"
    table.cell(0, 1).text = "Cell 2"
    table.cell(1, 0).text = "Cell 3"
    table.cell(1, 1).text = "Cell 4"
    buf = BytesIO()
    doc.save(buf)
    result = extract_from_docx(buf.getvalue())
    assert "Before table" in result
    assert "Cell 1" in result
    assert "Cell 4" in result
