from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from docxtpl import DocxTemplate, Subdoc

from app.html_to_docx import (
    RICH_TEXT_FIELDS,
    _extract_base64_images,
    enrich_context,
    enrich_context_with_subdoc,
    html_to_subdoc,
    html_to_text,
    is_html,
)


@pytest.fixture
def tpl(tmp_path: Path) -> DocxTemplate:
    d = Document()
    d.add_paragraph("{{p field }}")
    path = tmp_path / "t.docx"
    d.save(str(path))
    return DocxTemplate(str(path))


# --- is_html ---

def test_is_html_detects_tags():
    assert is_html("<p>hello</p>")
    assert is_html("<b>bold</b>")


def test_is_html_plain_text():
    assert not is_html("plain text")
    assert not is_html(None)
    assert not is_html("")


# --- html_to_text ---

def test_html_to_text_plain_passthrough():
    assert html_to_text("plain text") == "plain text"
    assert html_to_text(None) == ""
    assert html_to_text("") == ""


def test_html_to_text_strips_tags():
    assert html_to_text("<p>Hello <b>world</b></p>") == "Hello world"


def test_html_to_text_list_bullets():
    result = html_to_text("<ul><li>item1</li><li>item2</li></ul>")
    assert "• item1" in result
    assert "• item2" in result


def test_html_to_text_image_placeholder():
    png_b64 = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwADhQGAWjR9awAAAABJRU5ErkJggg=="
    result = html_to_text(f'<p>text</p><img src="data:image/png;base64,{png_b64}">')
    assert "[изображение]" in result


# --- html_to_subdoc ---

def test_subdoc_plain_text_as_str(tpl):
    assert html_to_subdoc("plain text", tpl) == "plain text"


def test_subdoc_empty(tpl):
    assert html_to_subdoc(None, tpl) == ""
    assert html_to_subdoc("", tpl) == ""


def test_subdoc_html_returns_subdoc(tpl):
    result = html_to_subdoc("<p>Hello <b>world</b></p>", tpl)
    assert isinstance(result, Subdoc)


def test_subdoc_renders_with_formatting(tmp_path: Path):
    """{{p field }} корректно рендерит Subdoc с форматированием."""
    d = Document()
    d.add_paragraph("before")
    d.add_paragraph("{{p content }}")
    d.add_paragraph("after")
    path = tmp_path / "tpl.docx"
    d.save(str(path))

    t = DocxTemplate(str(path))
    subdoc = html_to_subdoc("<p>Hello <b>World</b></p>", t)
    t.render({"content": subdoc})

    out = BytesIO()
    t.save(out)
    out.seek(0)

    texts = [p.text for p in Document(out).paragraphs]
    assert any("Hello" in t for t in texts)
    assert "before" in texts
    assert "after" in texts


# --- enrich_context (plain text only, no Subdoc) ---

def test_enrich_context_html_to_plain_text(tpl):
    ctx = {"bug_description": "<p>desc</p>", "bug_name": "SQLi"}
    result = enrich_context(ctx, tpl)
    assert result["bug_description"] == "desc"
    assert "bug_description_doc" not in result   # Subdoc не создаётся
    assert result["bug_name"] == "SQLi"


def test_enrich_context_plain_text_unchanged(tpl):
    ctx = {"reproduction_steps": "plain steps"}
    result = enrich_context(ctx, tpl)
    assert result["reproduction_steps"] == "plain steps"
    assert "reproduction_steps_doc" not in result


def test_enrich_context_none_unchanged(tpl):
    ctx = {"goal": None}
    result = enrich_context(ctx, tpl)
    assert result["goal"] is None


def test_rich_text_fields_set(tpl):
    assert RICH_TEXT_FIELDS == {
        "goal", "testConditions",
        "bug_description", "reproduction_steps", "remediation",
    }


# --- enrich_context_with_subdoc (opt-in Subdoc) ---

def test_enrich_context_with_subdoc_creates_doc(tpl):
    ctx = {"bug_description": "<p>desc</p>", "bug_name": "SQLi"}
    result = enrich_context_with_subdoc(ctx, tpl)
    assert result["bug_description"] == "desc"          # plain text сохраняется
    assert isinstance(result["bug_description_doc"], Subdoc)  # Subdoc добавлен
    assert result["bug_name"] == "SQLi"
    assert "bug_name_doc" not in result


def test_enrich_context_with_subdoc_plain_passthrough(tpl):
    ctx = {"reproduction_steps": "plain steps"}
    result = enrich_context_with_subdoc(ctx, tpl)
    assert result["reproduction_steps"] == "plain steps"
    assert result["reproduction_steps_doc"] == "plain steps"  # plain text, не Subdoc


# --- _extract_base64_images ---

PNG_B64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJ"
    "AAAADUlEQVR42mP8z8BQDwADhQGAWjR9awAAAABJRU5ErkJggg=="
)


def test_extract_base64_images_replaces_src():
    html = f'<img src="data:image/png;base64,{PNG_B64}">'
    result, tmp_files = _extract_base64_images(html)
    assert len(tmp_files) == 1
    assert tmp_files[0].exists()
    assert tmp_files[0].suffix == ".png"
    assert "data:" not in result
    for f in tmp_files:
        f.unlink(missing_ok=True)


def test_extract_no_data_uri():
    html = '<img src="http://example.com/img.png">'
    result, tmp_files = _extract_base64_images(html)
    assert len(tmp_files) == 0


def test_subdoc_with_base64_image(tpl):
    html = f'<p>text</p><img src="data:image/png;base64,{PNG_B64}">'
    result = html_to_subdoc(html, tpl)
    assert isinstance(result, Subdoc)
