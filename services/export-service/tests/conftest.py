import pytest
from pathlib import Path
from io import BytesIO
from docx import Document


@pytest.fixture
def simple_template(tmp_path: Path) -> Path:
    """Создаёт минимальный docx-шаблон с плейсхолдером."""
    doc = Document()
    doc.add_paragraph("{{ title }}")
    path = tmp_path / "template.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def vuln_template(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("{{ bug_name }} / {{ bug_criticality }}")
    path = tmp_path / "05_vulnerability.docx"
    doc.save(str(path))
    return path
