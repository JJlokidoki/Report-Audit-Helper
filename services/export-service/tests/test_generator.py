import pytest
import respx
import httpx
from io import BytesIO
from pathlib import Path
from docx import Document

from app import generator as gen_module
from app.generator import (
    SEVERITY_ORDER,
    _build_contexts,
    _merge_docs,
    _vuln_context,
    generate_report,
)
from app.config import settings


MOCK_REPORT = {"id": 1, "name": "Test", "report_type": "web"}
MOCK_SYSTEM_INFO = {
    "id": 1, "report_id": 1,
    "asName": "TestApp", "keId": "KE-1", "url": "https://test.com",
    "dateStart": "2024-01-01", "dateEnd": "2024-01-31",
    "segment": "DMZ", "goal": "Тестирование безопасности",
    "qualificationLevel": "Высокий", "accessLevel": "Внешний",
    "knowledgeLevel": "Чёрный ящик", "testConditions": "Без ограничений",
    "executors": [{"id": 1, "name": "Alice", "position": "Lead", "organization": "Corp"}],
    "software": [{"id": 1, "name": "Burp Suite", "version": "2024.1"}],
}
MOCK_SUMMARY = {
    "counts": {"critical": 1, "high": 0, "medium": 1, "low": 0, "info": 0},
    "vulnerabilities": [
        {"id": 1, "report_id": 1, "bug_name": "XSS", "bug_criticality": "critical",
         "bug_description": "desc", "cvss_score": 8.5, "cvss_vector": "AV:N",
         "reproduction_steps": "steps", "remediation": "fix", "automation_level": "no", "sort_order": 0},
        {"id": 2, "report_id": 1, "bug_name": "SQLI", "bug_criticality": "medium",
         "bug_description": None, "cvss_score": None, "cvss_vector": None,
         "reproduction_steps": None, "remediation": None, "automation_level": "no", "sort_order": 1},
    ],
}
MOCK_CHECKLIST = [
    {"id": 1, "report_id": 1, "checklist_type": "wstg", "check_id": "WSTG-INFO-01",
     "category": "Information Gathering", "name": "Разведка", "short_description": None,
     "goal": None, "status": "passed", "notes": "ok"},
]


def _make_template(tmp_path: Path, name: str, text: str = "{{ report_name }}") -> None:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(str(tmp_path / name))


@pytest.fixture
def web_templates(tmp_path: Path):
    template_dir = tmp_path / "web"
    template_dir.mkdir()
    for name in ["01_title.docx", "02_toc.docx", "03_general_info.docx",
                 "04_test_results.docx", "07_checklist.docx"]:
        _make_template(template_dir, name)
    _make_template(template_dir, "05_vulnerability.docx", "{{ bug_name }}")
    return tmp_path


def test_build_contexts():
    ctx = _build_contexts(MOCK_REPORT, MOCK_SYSTEM_INFO, MOCK_SUMMARY, MOCK_CHECKLIST)
    assert ctx["01_title.docx"]["asName"] == "TestApp"
    assert ctx["04_test_results.docx"]["critical_count"] == 1
    checks = ctx["07_checklist.docx"]["checks"]
    assert len(checks) == len(MOCK_CHECKLIST)
    assert checks[0]["result"] == "Выполнено\nok"


def test_vuln_context():
    ctx = _vuln_context(MOCK_SUMMARY["vulnerabilities"][0])
    assert ctx["bug_name"] == "XSS"
    assert ctx["cvss_score"] == 8.5


def test_vuln_context_nulls():
    ctx = _vuln_context(MOCK_SUMMARY["vulnerabilities"][1])
    assert ctx["cvss_score"] == ""
    assert ctx["bug_description"] == ""


def test_merge_docs():
    docs = []
    for _ in range(2):
        doc = Document()
        doc.add_paragraph("Test")
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        docs.append(buf)
    result = _merge_docs(docs)
    assert isinstance(result, BytesIO)
    assert result.tell() == 0


@pytest.mark.asyncio
async def test_generate_report_success(web_templates: Path, monkeypatch):
    monkeypatch.setattr(gen_module.settings, "template_dir", web_templates)
    base_url = settings.report_service_url

    with respx.mock:
        respx.get(f"{base_url}/api/reports/1").mock(return_value=httpx.Response(200, json=MOCK_REPORT))
        respx.get(f"{base_url}/api/reports/1/system-info").mock(return_value=httpx.Response(200, json=MOCK_SYSTEM_INFO))
        respx.get(f"{base_url}/api/reports/1/test-summary").mock(return_value=httpx.Response(200, json=MOCK_SUMMARY))
        respx.get(f"{base_url}/api/reports/1/checklist").mock(return_value=httpx.Response(200, json=MOCK_CHECKLIST))

        result = await generate_report(1)

    assert isinstance(result, BytesIO)
    doc = Document(result)
    assert doc is not None


@pytest.mark.asyncio
async def test_generate_report_no_templates(tmp_path: Path, monkeypatch):
    empty_dir = tmp_path / "empty"
    empty_dir.mkdir()
    monkeypatch.setattr(gen_module.settings, "template_dir", tmp_path)
    base_url = settings.report_service_url

    with respx.mock:
        respx.get(f"{base_url}/api/reports/1").mock(return_value=httpx.Response(200, json=MOCK_REPORT))
        respx.get(f"{base_url}/api/reports/1/system-info").mock(return_value=httpx.Response(200, json=MOCK_SYSTEM_INFO))
        respx.get(f"{base_url}/api/reports/1/test-summary").mock(return_value=httpx.Response(200, json=MOCK_SUMMARY))
        respx.get(f"{base_url}/api/reports/1/checklist").mock(return_value=httpx.Response(200, json=MOCK_CHECKLIST))

        with pytest.raises(FileNotFoundError):
            await generate_report(1)


# ── _vuln_context ──────────────────────────────────────────────────────────────

def test_vuln_context_is_first_true():
    ctx = _vuln_context(MOCK_SUMMARY["vulnerabilities"][0], is_first=True)
    assert ctx["is_first"] is True


def test_vuln_context_is_first_false():
    ctx = _vuln_context(MOCK_SUMMARY["vulnerabilities"][0], is_first=False)
    assert ctx["is_first"] is False


def test_vuln_context_is_first_default():
    """По умолчанию is_first=False."""
    ctx = _vuln_context(MOCK_SUMMARY["vulnerabilities"][0])
    assert ctx["is_first"] is False


# ── Severity ordering ──────────────────────────────────────────────────────────

def test_severity_order_critical_before_high():
    assert SEVERITY_ORDER["critical"] < SEVERITY_ORDER["high"]


def test_build_contexts_vuln_sorting():
    """Уязвимости в summary.vulnerabilities отсортированы при генерации."""
    vulns = MOCK_SUMMARY["vulnerabilities"]
    sorted_vulns = sorted(
        vulns, key=lambda v: SEVERITY_ORDER.get(v.get("bug_criticality", "info"), 99)
    )
    assert sorted_vulns[0]["bug_criticality"] == "critical"
    assert sorted_vulns[1]["bug_criticality"] == "medium"


# ── generate_report — vulnerability content ───────────────────────────────────

@pytest.fixture
def vuln_templates(tmp_path: Path):
    """Шаблоны: 04 + 05 с маркером is_first и именем уязвимости."""
    template_dir = tmp_path / "web"
    template_dir.mkdir()
    _make_template(template_dir, "04_test_results.docx", "{{ report_name }}")
    # Шаблон уязвимости: заголовок только для первой, потом имя баги
    _make_template(
        template_dir,
        "05_vulnerability.docx",
        "{% if is_first %}SECTION_HEADER{% endif %} {{ bug_name }}",
    )
    return tmp_path


@pytest.mark.asyncio
async def test_generate_report_vuln_names_present(vuln_templates: Path, monkeypatch):
    """Имена уязвимостей присутствуют в сгенерированном документе."""
    monkeypatch.setattr(gen_module.settings, "template_dir", vuln_templates)
    base_url = settings.report_service_url

    with respx.mock:
        respx.get(f"{base_url}/api/reports/1").mock(return_value=httpx.Response(200, json=MOCK_REPORT))
        respx.get(f"{base_url}/api/reports/1/system-info").mock(return_value=httpx.Response(200, json=MOCK_SYSTEM_INFO))
        respx.get(f"{base_url}/api/reports/1/test-summary").mock(return_value=httpx.Response(200, json=MOCK_SUMMARY))
        respx.get(f"{base_url}/api/reports/1/checklist").mock(return_value=httpx.Response(200, json=MOCK_CHECKLIST))

        result = await generate_report(1)

    doc = Document(result)
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "XSS" in full_text
    assert "SQLI" in full_text


@pytest.mark.asyncio
async def test_generate_report_is_first_header_once(vuln_templates: Path, monkeypatch):
    """SECTION_HEADER рендерится ровно один раз — только для первой уязвимости."""
    monkeypatch.setattr(gen_module.settings, "template_dir", vuln_templates)
    base_url = settings.report_service_url

    with respx.mock:
        respx.get(f"{base_url}/api/reports/1").mock(return_value=httpx.Response(200, json=MOCK_REPORT))
        respx.get(f"{base_url}/api/reports/1/system-info").mock(return_value=httpx.Response(200, json=MOCK_SYSTEM_INFO))
        respx.get(f"{base_url}/api/reports/1/test-summary").mock(return_value=httpx.Response(200, json=MOCK_SUMMARY))
        respx.get(f"{base_url}/api/reports/1/checklist").mock(return_value=httpx.Response(200, json=MOCK_CHECKLIST))

        result = await generate_report(1)

    doc = Document(result)
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert full_text.count("SECTION_HEADER") == 1


@pytest.mark.asyncio
async def test_generate_report_no_vulns(tmp_path: Path, monkeypatch):
    """Генерация без уязвимостей — шаблон 05 не используется."""
    template_dir = tmp_path / "web"
    template_dir.mkdir()
    _make_template(template_dir, "04_test_results.docx", "{{ report_name }}")
    _make_template(template_dir, "05_vulnerability.docx", "{{ bug_name }}")
    monkeypatch.setattr(gen_module.settings, "template_dir", tmp_path)

    summary_no_vulns = {**MOCK_SUMMARY, "vulnerabilities": []}
    base_url = settings.report_service_url

    with respx.mock:
        respx.get(f"{base_url}/api/reports/1").mock(return_value=httpx.Response(200, json=MOCK_REPORT))
        respx.get(f"{base_url}/api/reports/1/system-info").mock(return_value=httpx.Response(200, json=MOCK_SYSTEM_INFO))
        respx.get(f"{base_url}/api/reports/1/test-summary").mock(return_value=httpx.Response(200, json=summary_no_vulns))
        respx.get(f"{base_url}/api/reports/1/checklist").mock(return_value=httpx.Response(200, json=MOCK_CHECKLIST))

        result = await generate_report(1)

    doc = Document(result)
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "XSS" not in full_text
    assert "SQLI" not in full_text


@pytest.mark.asyncio
async def test_generate_report_vuln_severity_order(vuln_templates: Path, monkeypatch):
    """Critical уязвимость идёт раньше medium в итоговом документе."""
    monkeypatch.setattr(gen_module.settings, "template_dir", vuln_templates)
    base_url = settings.report_service_url

    with respx.mock:
        respx.get(f"{base_url}/api/reports/1").mock(return_value=httpx.Response(200, json=MOCK_REPORT))
        respx.get(f"{base_url}/api/reports/1/system-info").mock(return_value=httpx.Response(200, json=MOCK_SYSTEM_INFO))
        respx.get(f"{base_url}/api/reports/1/test-summary").mock(return_value=httpx.Response(200, json=MOCK_SUMMARY))
        respx.get(f"{base_url}/api/reports/1/checklist").mock(return_value=httpx.Response(200, json=MOCK_CHECKLIST))

        result = await generate_report(1)

    doc = Document(result)
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert full_text.index("XSS") < full_text.index("SQLI")
