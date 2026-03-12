from io import BytesIO
from pathlib import Path

from docx import Document

from app.filler import fill_template
from app.html_to_docx import RICH_TEXT_FIELDS, enrich_context_with_subdoc


def test_fill_template_replaces_placeholder(simple_template: Path):
    result = fill_template(simple_template, {"title": "Test Report"})
    assert isinstance(result, BytesIO)
    doc = Document(result)
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "Test Report" in full_text


def test_fill_template_empty_context(simple_template: Path):
    result = fill_template(simple_template, {})
    assert isinstance(result, BytesIO)
    doc = Document(result)
    assert doc is not None


def test_fill_template_returns_seeked_buffer(simple_template: Path):
    result = fill_template(simple_template, {"title": "X"})
    assert result.tell() == 0


def test_fill_template_html_plain_text(tmp_path: Path):
    """{{ field }} рендерит HTML как plain text."""
    d = Document()
    d.add_paragraph("{{ bug_description }}")
    path = tmp_path / "vuln.docx"
    d.save(str(path))

    result = fill_template(path, {"bug_description": "<p>Hello <b>World</b></p>"})
    assert isinstance(result, BytesIO)
    doc = Document(result)
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "Hello" in full_text


def test_fill_template_html_subdoc(tmp_path: Path):
    """{{p field_doc }} рендерит HTML как Subdoc — через enrich_context_with_subdoc."""
    from docxtpl import DocxTemplate

    d = Document()
    d.add_paragraph("{{p bug_description_doc }}")
    path = tmp_path / "vuln_rich.docx"
    d.save(str(path))

    tpl = DocxTemplate(str(path))
    ctx = enrich_context_with_subdoc({"bug_description": "<p>Hello <b>World</b></p>"}, tpl)
    tpl.render(ctx)
    out = BytesIO()
    tpl.save(out)
    out.seek(0)

    doc = Document(out)
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "Hello" in full_text


def test_rich_text_fields_set_is_correct():
    assert "bug_description" in RICH_TEXT_FIELDS
    assert "reproduction_steps" in RICH_TEXT_FIELDS
    assert "remediation" in RICH_TEXT_FIELDS
    assert "goal" in RICH_TEXT_FIELDS
