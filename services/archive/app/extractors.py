from io import BytesIO

import httpx
from docx import Document


def extract_from_docx(file_bytes: bytes) -> str:
    doc = Document(BytesIO(file_bytes))
    parts: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def extract_from_pdf(file_bytes: bytes) -> str:
    import pymupdf

    doc = pymupdf.open(stream=file_bytes, filetype="pdf")
    pages: list[str] = []
    for page in doc:
        text = page.get_text().strip()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


async def extract_from_report_service(report_id: int, base_url: str) -> tuple[str, dict]:
    async with httpx.AsyncClient(timeout=30, trust_env=False) as client:
        report_resp = await client.get(f"{base_url}/api/reports/{report_id}")
        report_resp.raise_for_status()
        report = report_resp.json()

        si_resp = await client.get(f"{base_url}/api/reports/{report_id}/system-info")
        si_resp.raise_for_status()
        system_info = si_resp.json()

        vulns_resp = await client.get(f"{base_url}/api/reports/{report_id}/vulnerabilities")
        vulns_resp.raise_for_status()
        vulns = vulns_resp.json()

    sections: list[str] = []
    sections.append(f"# Отчёт: {report.get('name', '')}")
    sections.append(f"Тип: {report.get('report_type', '')}")

    if system_info.get("asName"):
        sections.append(f"Система: {system_info['asName']}")
    if system_info.get("url"):
        sections.append(f"URL: {system_info['url']}")
    if system_info.get("description"):
        sections.append(f"Описание: {system_info['description']}")

    software = system_info.get("software", [])
    if software:
        sw_names = ", ".join(s["name"] for s in software)
        sections.append(f"Технологии: {sw_names}")

    for vuln in vulns:
        sections.append(f"\n## {vuln.get('bug_name', 'Без названия')}")
        sections.append(f"Критичность: {vuln.get('bug_criticality', '')}")
        if vuln.get("cvss_score"):
            sections.append(f"CVSS: {vuln['cvss_score']}")
        if vuln.get("bug_description"):
            sections.append(f"Описание: {vuln['bug_description']}")
        if vuln.get("reproduction_steps"):
            sections.append(f"Шаги воспроизведения: {vuln['reproduction_steps']}")
        if vuln.get("remediation"):
            sections.append(f"Рекомендации: {vuln['remediation']}")

    metadata = {
        "report_type": report.get("report_type", ""),
        "technology": ", ".join(s["name"] for s in software) if software else "",
    }

    return "\n".join(sections), metadata
